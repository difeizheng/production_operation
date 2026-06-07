"""Step 2: V4.docx 原因段落盘点 - 识别需要映射到 Excel 的"原因/分析"段落

运行方式：
    PYTHONPATH=. python examples/inventory_v4_reasons.py

输出：
    - 终端打印所有段落及分类
    - JSON 文件保存"待映射清单"（reason_segments.json）

识别策略：
    1. 关键词匹配："原因"、"原因为"、"主要受...影响"、"受...拉动"等
    2. 结构特征：含 "1、" "2、" "一是" "二是" 等
    3. 段落长度：> 50 字且非纯数据
    4. 跳过标题/空段/纯表格标记段

分类：
    - reason: 原因/归因段落
    - data: 纯数据段落
    - summary: 概述段（含数据但也有原因）
    - heading: 标题段
    - table_ref: 表格引用段
    - skip: 跳过（标题/空）
"""
from __future__ import annotations

import io
import json
import logging
import re
import sys
from dataclasses import dataclass, field, asdict
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))


# ============================================================================
# 数据类
# ============================================================================

@dataclass(frozen=True)
class Segment:
    """V4 中的一个段落。"""
    index: int                 # 段落编号 P0, P1, ...
    style: str                 # Normal / Heading 1 / Heading 2 / Heading 3
    text: str                  # 段落文本
    category: str              # reason / data / summary / heading / table_ref / skip
    needs_mapping: bool        # True = 需要映射到 Excel
    reason_keywords: List[str] = field(default_factory=list)  # 命中的关键词
    priority: int = 0          # 映射优先级 0-3（数字越大越优先）
    suggested_slot: Optional[str] = None  # 建议的 ReasonSlot ID


# ============================================================================
# 关键词与分类规则
# ============================================================================

# 原因/分析关键词
REASON_KEYWORDS = [
    r"主要原因为?", r"主要原因", r"原因为?", r"原因是", r"原因为：",
    r"主要受.+?影响", r"受.+?影响", r"受.+?拉动", r"受.+?下降",
    r"主要是", r"主因", r"其中", r"一是.+?二是.+?三是",
    r"^1、.+?2、", r"^1\..+?2\.",
]

# 表格引用
TABLE_REF_PATTERNS = [
    r"^表\s*\d+", r"^附表",
]

# 数据段落（纯数字）
DATA_ONLY_PATTERN = re.compile(r"^[\d\.\,\s年月日万亿千瓦时元分百分点%]+$")

# 周/年关键词
WEEKLY_KEYWORDS = [r"上周", r"本周", r"第\d+周", r"环比", r"同比", r"较去年同期"]


# ============================================================================
# 段落分类器
# ============================================================================

class SegmentClassifier:
    """段落分类器：识别"原因/分析"段落。"""

    def __init__(self) -> None:
        self._reason_patterns = [re.compile(p) for p in REASON_KEYWORDS]

    def classify(self, index: int, style: str, text: str) -> Segment:
        """分类单个段落。"""
        text = text.strip()
        if not text:
            return Segment(
                index=index, style=style, text="",
                category="skip", needs_mapping=False,
            )

        # 标题
        if "Heading" in style or "Title" in style:
            return Segment(
                index=index, style=style, text=text,
                category="heading", needs_mapping=False,
            )

        # 表格引用
        for pattern in TABLE_REF_PATTERNS:
            if re.match(pattern, text):
                return Segment(
                    index=index, style=style, text=text,
                    category="table_ref", needs_mapping=False,
                )

        # 匹配原因关键词
        matched_keywords = self._match_reason_keywords(text)

        # 决定分类
        has_data = self._is_data_heavy(text)
        has_reason = len(matched_keywords) > 0

        if has_reason and has_data:
            category = "summary"   # 数据+原因混合
        elif has_reason:
            category = "reason"    # 纯原因
        elif has_data and len(text) < 200:
            category = "data"      # 纯数据
        else:
            category = "other"     # 其他叙述

        needs_mapping = category in ("reason", "summary")
        priority = self._calc_priority(category, matched_keywords, text)

        return Segment(
            index=index, style=style, text=text,
            category=category, needs_mapping=needs_mapping,
            reason_keywords=matched_keywords,
            priority=priority,
            suggested_slot=self._suggest_slot(text, category),
        )

    def _match_reason_keywords(self, text: str) -> List[str]:
        """找出所有命中的原因关键词。"""
        matched = []
        for pattern in self._reason_patterns:
            m = pattern.search(text)
            if m:
                matched.append(m.group(0))
        return matched

    def _is_data_heavy(self, text: str) -> bool:
        """判断段落是否数据密集（数字占比 > 20%）。"""
        digits = sum(1 for c in text if c.isdigit() or c in ".")
        return digits / max(len(text), 1) > 0.2

    def _calc_priority(
        self, category: str, keywords: List[str], text: str
    ) -> int:
        """计算映射优先级（0-3）。"""
        if category == "skip":
            return 0
        if category == "reason" and len(keywords) >= 2:
            return 3   # 多关键词 → 高优先级
        if category == "reason":
            return 2
        if category == "summary":
            return 2
        return 1

    def _suggest_slot(self, text: str, category: str) -> Optional[str]:
        """根据文本内容建议对应的 ReasonSlot ID。"""
        if not text:
            return None

        # 简单的关键词匹配（实际映射时人工确认）
        suggestions = []

        # 国内/国际 标识
        is_intl = "国际" in text
        is_dom = "国内" in text or "集团" in text

        # 同比/环比
        is_yoy = "同比" in text
        is_wow = "环比" in text

        # 电量/电价
        is_elec = "电量" in text
        is_price = "电价" in text or "价格" in text

        # 品类
        category_keywords = {
            "hydro": ["水电", "金沙江", "溪洛渡", "向家坝", "乌东德", "白鹤滩", "三峡"],
            "wind": ["风电"],
            "solar": ["光伏"],
            "thermal": ["火电", "煤电"],
            "new_energy": ["新能源"],
        }

        if is_intl:
            if is_yoy:
                return "intl.yoy.t" if "T" in str(self) else "intl.yoy.h"
            if is_wow:
                return "intl.wow.t" if "T" in str(self) else "intl.wow.h"

        if is_dom:
            prefix = "dom"
            period = "yoy" if is_yoy else ("wow" if is_wow else "")
            metric = "elec" if is_elec else ("price" if is_price else "")

            if period and metric:
                # 检查品类
                for cat_name, cat_keywords in category_keywords.items():
                    if any(kw in text for kw in cat_keywords):
                        return f"{prefix}.{metric}.{period}.{cat_name}"
                # 默认 长江电力+全集团
                return f"{prefix}.{metric}.{period}.changjiang"

        return None


# ============================================================================
# 主流程
# ============================================================================

def inventory_v4(v4_path: Path) -> List[Segment]:
    """盘点 V4.docx 全部段落。"""
    logger.info("加载 V4.docx: %s", v4_path)
    doc = Document(str(v4_path))
    logger.info("段落数: %d", len(doc.paragraphs))

    classifier = SegmentClassifier()
    segments: List[Segment] = []

    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        style = p.style.name if p.style else "Normal"
        seg = classifier.classify(i, style, text)
        segments.append(seg)

    return segments


def print_inventory(segments: List[Segment]) -> None:
    """打印盘点结果。"""
    print("\n" + "=" * 80)
    print("V4.docx 段落盘点结果")
    print("=" * 80)

    # 按类别统计
    by_category: Dict[str, int] = {}
    needs_mapping_count = 0
    for seg in segments:
        by_category[seg.category] = by_category.get(seg.category, 0) + 1
        if seg.needs_mapping:
            needs_mapping_count += 1

    print(f"\n总计: {len(segments)} 段")
    print(f"需映射: {needs_mapping_count} 段")
    print(f"\n按类别:")
    for cat, count in sorted(by_category.items()):
        print(f"  {cat}: {count}")

    print("\n" + "-" * 80)
    print("【需映射段落详情】")
    print("-" * 80)

    for seg in segments:
        if not seg.needs_mapping:
            continue
        marker = "*" * seg.priority
        print(f"\nP{seg.index} {marker} [{seg.category}] 建议: {seg.suggested_slot or '(无)'}")
        print(f"  关键词: {seg.reason_keywords[:3]}")
        print(f"  文本: {seg.text[:200]}{'...' if len(seg.text) > 200 else ''}")


def export_inventory(segments: List[Segment], output_path: Path) -> None:
    """导出待映射清单为 JSON。"""
    needs_mapping = [asdict(s) for s in segments if s.needs_mapping]

    data = {
        "source": "详版_第22周例会营销发言材料_V4.docx",
        "total_segments": len(segments),
        "needs_mapping_count": len(needs_mapping),
        "needs_mapping": needs_mapping,
    }

    output_path.parent.mkdir(parents=True, exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

    print(f"\n✅ 待映射清单已保存: {output_path}")
    print(f"   共 {len(needs_mapping)} 段需映射")


def main() -> None:
    # Windows 控制台 UTF-8 支持（仅运行时）
    if sys.platform == "win32":
        try:
            sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
            sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")
        except Exception:
            pass  # pytest 环境可能没有 buffer

    # 配置日志（仅在主入口时启用）
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s [%(levelname)s] %(message)s",
    )

    v4_path = PROJECT_ROOT / "files" / "详版_第22周例会营销发言材料_V4.docx"
    output_path = PROJECT_ROOT / "data" / "processed" / "reason_segments.json"

    if not v4_path.exists():
        print(f"❌ V4 文件不存在: {v4_path}")
        return

    segments = inventory_v4(v4_path)
    print_inventory(segments)
    export_inventory(segments, output_path)

    print("\n" + "=" * 80)
    print("📌 下一步: Step 3 - 批量标注 reason_map.json")
    print("=" * 80)
    print("运行: PYTHONPATH=. python examples/build_reason_map.py")


if __name__ == "__main__":
    main()
