"""销售表 subdoc 改造脚本 v2 - 简化方案

策略（不依赖 docxtpl 特殊语法）：
    1. 保留 V3 模板的 10 行结构
    2. 每行的 6 个数据单元格改为 {{ row_<idx>.<col> }} 形式
    3. 报告生成器构造 9 个 row_<idx> 字典，每个含 6 个列键

示例：
    第 2 行（水电行）：{{ row_1.hydro }} | {{ row_1.new_energy }} | ...
    第 3 行（同比行）：{{ row_2.hydro }} | {{ row_2.new_energy }} | ...
    ...
    第 10 行（收入环比）：{{ row_9.hydro }} | {{ row_9.new_energy }} | ...

优势：
    - 不需要 docxtpl 特殊语法（兼容所有 Jinja2 实现）
    - 测试简单（每个 row 是独立 dict）
    - 数据来源清晰（构造 9 个 dict 即可）
"""
from __future__ import annotations

import io
import logging
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional

PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

logger = logging.getLogger(__name__)


# ============================================================================
# 销售表行定义
# ============================================================================

# 销售表行定义：每行的标签 + data key
SALES_ROW_DEFS = [
    {"label": "国内上网电量", "data_key": "electricity", "format": "number"},
    {"label": "同比", "data_key": "yoy_electricity", "format": "pct"},
    {"label": "环比", "data_key": "wow_electricity", "format": "pct"},
    {"label": "国内上网电价", "data_key": "price", "format": "number"},
    {"label": "电价同比", "data_key": "yoy_price", "format": "fen"},
    {"label": "电价环比", "data_key": "wow_price", "format": "fen"},
    {"label": "国内发电收入", "data_key": "revenue", "format": "number"},
    {"label": "收入同比", "data_key": "yoy_revenue", "format": "pct"},
    {"label": "收入环比", "data_key": "wow_revenue", "format": "pct"},
]

# 销售表列定义（5 品类 + 合计 = 6 列数据）
SALES_COL_DEFS = [
    {"key": "hydro"},
    {"key": "new_energy"},
    {"key": "wind"},
    {"key": "solar"},
    {"key": "thermal"},
    {"key": "total"},
]


# ============================================================================
# 模板改造函数
# ============================================================================

def _set_cell_text(cell: Any, text: str) -> None:
    """设置单元格文本（保留第一个 run 的格式）。"""
    if cell.paragraphs:
        first_para = cell.paragraphs[0]
        if first_para.runs:
            first_run = first_para.runs[0]
            first_run.text = text
            for run in first_para.runs[1:]:
                run.text = ""
        else:
            first_para.add_run(text)
    else:
        cell.add_paragraph(text)


def prepare_sales_table(
    template_path: Path = PROJECT_ROOT / "data" / "templates" / "report_template_jinja.docx",
) -> None:
    """把销售表（10×7）的数据单元格改为 row_<idx>.<col> 形式。"""
    from docx import Document

    if not template_path.exists():
        logger.error("模板不存在: %s", template_path)
        return

    doc = Document(str(template_path))

    # 找销售表（10 行 × 7 列）
    sales_table = None
    for table in doc.tables:
        if len(table.rows) == 10 and len(table.columns) == 7:
            sales_table = table
            break

    if sales_table is None:
        logger.warning("未找到销售表（10×7）")
        return

    logger.info("找到销售表（10×7）")

    # 改造数据行（行 1-9）
    for row_idx in range(1, 10):
        row = sales_table.rows[row_idx]
        if len(row.cells) < 7:
            continue

        # 标签列（col 0）保留原值

        # 数据列（col 1-6）替换为 {{ row_<row_idx>.<col_key> }}
        for col_idx in range(1, 7):
            cell = row.cells[col_idx]
            col_def = SALES_COL_DEFS[col_idx - 1]
            placeholder = f"{{{{ row_{row_idx}.{col_def['key']} }}}}"
            _set_cell_text(cell, placeholder)

    doc.save(str(template_path))
    logger.info("✅ 销售表已改造（9 数据行 × 6 数据列，row_1 ~ row_9）")


def prepare_spot_table(
    template_path: Path = PROJECT_ROOT / "data" / "templates" / "report_template_jinja.docx",
) -> None:
    """把现货表（4×11）的数据单元格改为 region.<key> 形式。

    现货表结构：
        Row 0: 地区（广东、山西、...）
        Row 1: 均价
        Row 2: 同比
        Row 3: 环比

    改造后：
        Row 1: {{ spot_avg.广东 }}, {{ spot_avg.山西 }}, ...
        Row 2: {{ spot_yoy.广东 }}, ...
        Row 3: {{ spot_wow.广东 }}, ...
    """
    from docx import Document

    if not template_path.exists():
        return

    doc = Document(str(template_path))

    # 找现货表（4 行 × 11 列）
    spot_table = None
    for table in doc.tables:
        if len(table.rows) == 4 and len(table.columns) == 11:
            spot_table = table
            break

    if spot_table is None:
        logger.warning("未找到现货表（4×11）")
        return

    logger.info("找到现货表（4×11）")

    # 提取地区名（行 0，col 1-10）
    region_names = []
    for col_idx in range(1, 11):
        cell = spot_table.rows[0].cells[col_idx]
        region_names.append(cell.text.strip())

    # 改造数据行
    metric_keys = ["spot_avg", "spot_yoy", "spot_wow"]
    for row_idx in [1, 2, 3]:
        row = spot_table.rows[row_idx]
        metric_key = metric_keys[row_idx - 1]
        for col_idx in range(1, 11):
            cell = row.cells[col_idx]
            region = region_names[col_idx - 1]
            placeholder = f"{{{{ {metric_key}.{region} }}}}"
            _set_cell_text(cell, placeholder)

    doc.save(str(template_path))
    logger.info(f"✅ 现货表已改造（地区: {region_names[:3]}...，{len(region_names)} 个）")


# ============================================================================
# 数据构造（供 report_generator_v2.py 调用）
# ============================================================================

def build_sales_rows(data: Dict[str, Any]) -> Dict[str, Dict[str, str]]:
    """构造销售表 9 行数据。

    优先从 report_table_1.data 读取（已转换单位的最终数据），
    否则从嵌套结构 domestic.electricity.hydro 等读取。

    Returns:
        {"row_1": {"hydro": "59.05", ...}, "row_2": {...}, ...}
    """
    # 优先使用 report_table_1（已转换为亿千瓦时/亿元）
    rt1 = data.get("report_table_1", {})
    rt1_data = rt1.get("data", [])

    # 备选：嵌套结构（domestic.electricity.hydro 等）
    domestic = data.get("domestic", {})

    # 映射：行号 → row_def
    row_defs_map = {idx + 1: defn for idx, defn in enumerate(SALES_ROW_DEFS)}

    result: Dict[str, Dict[str, str]] = {}

    for row_idx in range(1, 10):
        row_key = f"row_{row_idx}"
        row_def = row_defs_map.get(row_idx)
        if not row_def:
            continue

        row_data: Dict[str, str] = {}
        for col_idx, col_def in enumerate(SALES_COL_DEFS, start=1):
            # 三层 fallback 路径
            value: Any = None

            # 路径 1: report_table_1.data[row_idx-1][col_idx-1]
            if rt1_data and len(rt1_data) >= row_idx and len(rt1_data[row_idx - 1]) >= col_idx:
                value = rt1_data[row_idx - 1][col_idx - 1]

            # 路径 2: 嵌套结构 (domestic.electricity.hydro)
            if value is None:
                nested_key = f"{row_def['data_key']}.{col_def['key']}"
                parts = nested_key.split(".")
                cur: Any = domestic
                for p in parts:
                    if isinstance(cur, dict):
                        cur = cur.get(p)
                    else:
                        cur = None
                        break
                value = cur

            # 路径 3: flat report.* 键
            if value is None:
                flat_key = f"report.{row_def['data_key']}.{col_def['key']}"
                value = data.get(flat_key)

            row_data[col_def["key"]] = _format_value(value, row_def["format"])

        result[row_key] = row_data

    return result


def build_spot_rows(data: Dict[str, Any]) -> Dict[str, Dict[str, str]]:
    """构造现货表 3 行数据。

    Returns:
        {"spot_avg": {"广东": "0.366", "山西": "0.273", ...}, ...}
    """
    spot_prices = data.get("spot_prices", {}).get("data", {})

    result: Dict[str, Dict[str, str]] = {
        "spot_avg": {},
        "spot_yoy": {},
        "spot_wow": {},
    }

    for region, region_data in spot_prices.items():
        result["spot_avg"][region] = _format_value(region_data.get("avg"), "yuan_per_kwh")
        result["spot_yoy"][region] = _format_value(region_data.get("yoy"), "pct_decimal")
        result["spot_wow"][region] = _format_value(region_data.get("wow"), "pct_decimal")

    return result


def _format_value(value: Any, fmt: str) -> str:
    """格式化数据值。"""
    if value is None:
        return "—"

    if fmt == "pct":
        # 比率 → 百分比（带符号）
        pct = value * 100
        return f"{pct:+.1f}%" if pct != 0 else "0.0%"
    elif fmt == "pct_decimal":
        # 比率 → 百分比（带符号，小数）
        pct = value * 100
        return f"{pct:+.1f}%"
    elif fmt == "fen":
        # 元 → 分（带符号）
        fen = value * 100
        return f"{fen:+.1f}"
    elif fmt == "yuan_per_kwh":
        return f"{value:.3f}"
    else:  # "number"
        return f"{value:.2f}"


# ============================================================================
# 主流程
# ============================================================================

def main() -> None:
    if sys.platform == "win32":
        try:
            sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
        except Exception:
            pass

    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s [%(levelname)s] %(message)s",
    )

    prepare_sales_table()
    prepare_spot_table()


if __name__ == "__main__":
    main()
