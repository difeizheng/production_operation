"""Integration tests for report_generator + reason_resolver

测试覆盖：
    1. 旧模式（无 reason_text）保持向后兼容
    2. 新模式（带 reason_text）渲染原因文本
    3. 占位符 fallback 行为
    4. 端到端生成 Word 文档
"""
from __future__ import annotations

import json
import sys
from pathlib import Path
from types import MethodType
from typing import Any, Dict

import pytest

PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

from docx import Document

from src.collector.reason_collector import (
    ALL_REASON_SLOTS,
    ReasonCollector,
    ReasonResult,
    ReasonSlot,
)
from src.generator.reason_polisher import PolishResult, ReasonPolisher
from src.generator.reason_resolver import (
    ReasonResolver,
    ResolvedSegment,
)
from src.generator.report_generator import (
    ReportGenerator,
    _render_reason,
    _build_reason_paragraphs,
)


# ============================================================================
# Fixtures
# ============================================================================

@pytest.fixture
def sample_data() -> Dict[str, Any]:
    """最小可用的 data 字典。"""
    return {
        "meta": {"year": 2026, "week": 21, "start_date": "5月16日", "end_date": "5月22日"},
        "international": {"price": {"total": 0.32}, "yoy": {"price_change": 0.039}},
        "domestic": {
            "electricity": {"total": 80.3, "hydro": 53.5, "thermal": 11.2},
            "price": {"total": 0.311, "hydro": 0.288, "thermal": 0.382},
        },
        "yoy": {
            "electricity": {"total": 0.033},
            "price": {"total": -0.009},
        },
        "wow": {
            "electricity": {"total": 0.157},
            "price": {"total": 0.001},
        },
        "revenue": {"total": 249000},  # 万元
    }


@pytest.fixture
def sample_reason_text() -> Dict[str, ResolvedSegment]:
    """模拟 ReasonResolver 输出的字典。"""
    def make_seg(placeholder: str, text: str, level: str = "HIGH") -> ResolvedSegment:
        return ResolvedSegment(
            placeholder=placeholder,
            final_text=text,
            automation_level=level,
            source_slots=["test.slot"],
            raw_text=text,
            polished=True,
            is_fallback=False,
            tokens_used=50,
        )
    return {
        "{{ v4_P6_dom_elec_yoy_wow }}": make_seg(
            "{{ v4_P6_dom_elec_yoy_wow }}",
            "1、全口径电量同比增加，主要原因为梯级电站来水偏丰。",
        ),
        "{{ v4_P11_dom_price_yoy }}": make_seg(
            "{{ v4_P11_dom_price_yoy }}",
            "水电电价同比下降0.1分，主要原因是金下梯级电站占比下降。",
        ),
        "{{ v4_P20_intl_price_yoy }}": make_seg(
            "{{ v4_P20_intl_price_yoy }}",
            "国际电价同比提高3.9分，主要原因是拉美售电合同价格提高。",
        ),
        "{{ v4_P23_market_hydro }}": make_seg(
            "{{ v4_P23_market_hydro }}",
            "水电项目平均交易电价0.304元，同比提高1.3分。",
        ),
        # 故意不放 P13-P16 → 走 fallback
    }


# ============================================================================
# 辅助函数测试
# ============================================================================

class TestRenderReasonHelper:
    """_render_reason 辅助函数测试。"""

    def test_render_with_text(self) -> None:
        seg = ResolvedSegment(
            placeholder="{{ test }}", final_text="原因文本",
            automation_level="HIGH", source_slots=[], raw_text="",
            polished=True, is_fallback=False,
        )
        result = _render_reason({"{{ test }}": seg}, "{{ test }}")
        assert result == "原因文本"

    def test_render_empty_fallback(self) -> None:
        seg = ResolvedSegment(
            placeholder="{{ test }}", final_text="",
            automation_level="MANUAL", source_slots=[], raw_text="",
            polished=False, is_fallback=True,
        )
        result = _render_reason({"{{ test }}": seg}, "{{ test }}", "[默认]")
        assert result == "[默认]"

    def test_render_missing_placeholder(self) -> None:
        result = _render_reason(None, "{{ test }}", "[未提供]")
        assert result == "[未提供]"


class TestBuildReasonParagraphsHelper:
    """_build_reason_paragraphs 辅助函数测试。"""

    def test_no_reason_text(self) -> None:
        doc = Document()
        _build_reason_paragraphs(doc, None, ["{{ a }}", "{{ b }}"])
        # 不应添加任何段落
        assert len(doc.paragraphs) == 0

    def test_skip_empty_text(self) -> None:
        doc = Document()
        seg_empty = ResolvedSegment(
            placeholder="{{ a }}", final_text="",
            automation_level="MANUAL", source_slots=[], raw_text="",
            polished=False, is_fallback=True,
        )
        seg_filled = ResolvedSegment(
            placeholder="{{ b }}", final_text="有内容",
            automation_level="HIGH", source_slots=[], raw_text="",
            polished=True, is_fallback=False,
        )
        _build_reason_paragraphs(
            doc,
            {"{{ a }}": seg_empty, "{{ b }}": seg_filled},
            ["{{ a }}", "{{ b }}"],
        )
        # 只有 1 个段落被添加
        assert len(doc.paragraphs) == 1
        assert doc.paragraphs[0].text == "有内容"


# ============================================================================
# ReportGenerator 集成测试
# ============================================================================

class TestReportGeneratorIntegration:
    """ReportGenerator 集成测试。"""

    def test_generate_without_reason_text(
        self, sample_data: Dict[str, Any], tmp_path
    ) -> None:  # type: ignore[no-untyped-def]
        """无 reason_text 时使用 fallback（向后兼容）。"""
        generator = ReportGenerator(output_dir=str(tmp_path))
        output_path = generator.generate_report(
            data=sample_data, year=2026, week=21,
        )

        assert Path(output_path).exists()
        doc = Document(output_path)

        # 文档应包含（待补充）占位（因为 reason_text 是 None）
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "（待补充）" in all_text or "待补充" in all_text

    def test_generate_with_reason_text(
        self, sample_data: Dict[str, Any],
        sample_reason_text: Dict[str, ResolvedSegment],
        tmp_path,
    ) -> None:  # type: ignore[no-untyped-def]
        """带 reason_text 时渲染原因文本。"""
        generator = ReportGenerator(output_dir=str(tmp_path))
        output_path = generator.generate_report(
            data=sample_data, year=2026, week=21,
            reason_text=sample_reason_text,
        )

        assert Path(output_path).exists()
        doc = Document(output_path)
        all_text = "\n".join(p.text for p in doc.paragraphs)

        # 应该包含样例 reason_text 中的内容
        assert "梯级电站来水偏丰" in all_text, "P6 原因应被渲染"
        assert "金下梯级电站占比下降" in all_text, "P11 原因应被渲染"
        assert "拉美售电合同价格提高" in all_text, "P20 原因应被渲染"
        assert "水电项目平均交易电价" in all_text, "P23 原因应被渲染"

    def test_generate_does_not_break_on_empty_reason_text(
        self, sample_data: Dict[str, Any], tmp_path
    ) -> None:  # type: ignore[no-untyped-def]
        """reason_text 为空字典时不崩。"""
        generator = ReportGenerator(output_dir=str(tmp_path))
        output_path = generator.generate_report(
            data=sample_data, year=2026, week=21,
            reason_text={},
        )
        assert Path(output_path).exists()

    def test_partial_reason_text_renders_available(
        self, sample_data: Dict[str, Any], tmp_path
    ) -> None:  # type: ignore[no-untyped-def]
        """只有部分 reason_text 时，可用的渲染，缺失的走 fallback。"""
        partial_text = {
            "{{ v4_P6_dom_elec_yoy_wow }}": ResolvedSegment(
                placeholder="{{ v4_P6_dom_elec_yoy_wow }}",
                final_text="部分原因",
                automation_level="HIGH", source_slots=[], raw_text="",
                polished=True, is_fallback=False,
            ),
        }
        generator = ReportGenerator(output_dir=str(tmp_path))
        output_path = generator.generate_report(
            data=sample_data, year=2026, week=21,
            reason_text=partial_text,
        )

        doc = Document(output_path)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        # 部分内容应被渲染
        assert "部分原因" in all_text
