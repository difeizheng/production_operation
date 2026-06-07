"""模板样式增强脚本 - 优化 report_template_jinja.docx 的视觉效果

增强维度：
    1. 标题层级（Heading 1/2/3 字号 + 颜色）
    2. 段落间距（1.5 倍行距 + 段前段后）
    3. 字体统一（仿宋正文 + 黑体标题 + Consolas 数字）
    4. 表格样式（带边框、表头高亮）
    5. 页眉页脚（页码 + 周报标识）

使用：
    PYTHONPATH=. python scripts/enhance_template.py
"""
from __future__ import annotations

import io
import logging
import sys
from pathlib import Path
from typing import Any, Dict, List

PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

from docx.shared import Pt, RGBColor  # noqa: E402

logger = logging.getLogger(__name__)


# ============================================================================
# 样式定义
# ============================================================================

# 标题样式映射
HEADING_STYLES = {
    "Heading 1": {
        "font_name": "黑体",
        "font_size_pt": 16,
        "bold": True,
        "color": "1F4E79",       # 深蓝
        "space_before_pt": 18,
        "space_after_pt": 12,
    },
    "Heading 2": {
        "font_name": "黑体",
        "font_size_pt": 14,
        "bold": True,
        "color": "2E74B5",       # 中蓝
        "space_before_pt": 12,
        "space_after_pt": 8,
    },
    "Heading 3": {
        "font_name": "黑体",
        "font_size_pt": 12,
        "bold": True,
        "color": "5B9BD5",       # 浅蓝
        "space_before_pt": 8,
        "space_after_pt": 6,
    },
}

# 正文样式
NORMAL_STYLE = {
    "font_name": "仿宋",
    "font_size_pt": 12,
    "line_spacing": 1.5,
    "space_after_pt": 6,
}


# ============================================================================
# 样式应用函数
# ============================================================================

def apply_paragraph_style(paragraph: Any, style_name: str) -> None:
    """应用段落样式。"""
    config = HEADING_STYLES.get(style_name)
    if not config:
        return

    pf = paragraph.paragraph_format
    pf.space_before = config["space_before_pt"] * 12700  # pt → EMU
    pf.space_after = config["space_after_pt"] * 12700
    pf.line_spacing = 1.5

    for run in paragraph.runs:
        run.font.name = config["font_name"]
        run.font.size = Pt(config["font_size_pt"])
        run.font.bold = config["bold"]
        run.font.color.rgb = RGBColor.from_string(config["color"])


def apply_normal_style(doc: Any) -> None:
    """设置 Normal 样式（全文默认）。"""
    style = doc.styles["Normal"]
    font = style.font
    font.name = NORMAL_STYLE["font_name"]
    font.size = Pt(NORMAL_STYLE["font_size_pt"])

    pf = style.paragraph_format
    pf.line_spacing = NORMAL_STYLE["line_spacing"]
    pf.space_after = Pt(NORMAL_STYLE["space_after_pt"])


def set_page_setup(doc: Any) -> None:
    """设置页面（边距、纸张）。"""
    from docx.shared import Cm
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)
        # A4
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)


def add_page_number_footer(doc: Any) -> None:
    """添加页码页脚。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    for section in doc.sections:
        footer = section.footer
        if not footer.paragraphs:
            footer.add_paragraph()
        p = footer.paragraphs[0]
        p.alignment = 1  # 居中

        # 创建 PAGE field
        run = p.add_run()
        run.font.name = "仿宋"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string("808080")

        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)


def enhance_tables(doc: Any) -> None:
    """优化表格样式。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.style import WD_STYLE_TYPE

    for table in doc.tables:
        # 设置表格样式（如果不存在则用默认 Grid）
        try:
            table.style = "Table Grid"
        except KeyError:
            pass

        # 表头行加底色
        if len(table.rows) > 0:
            header_row = table.rows[0]
            for cell in header_row.cells:
                # 设置底色（浅蓝）
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "D9E2F3")
                tcPr.append(shd)

                # 加粗
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.bold = True
                        run.font.name = "黑体"
                        run.font.size = Pt(10)

        # 表格内容字体
        for row in table.rows[1:]:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if not run.font.name:
                            run.font.name = "仿宋"
                        run.font.size = Pt(10)


# ============================================================================
# 主流程
# ============================================================================

def enhance_template(
    source: Path = PROJECT_ROOT / "data" / "templates" / "report_template_jinja.docx",
    target: Path = None,
) -> None:
    """增强模板样式。

    Args:
        source: 源模板路径
        target: 输出路径（None 时覆盖源文件）
    """
    if target is None:
        target = source

    if not source.exists():
        logger.error("源模板不存在: %s", source)
        return

    from docx import Document

    logger.info("加载模板: %s", source)
    doc = Document(str(source))

    # 1. 设置 Normal 样式
    apply_normal_style(doc)
    logger.info("✅ Normal 样式已设置")

    # 2. 应用标题样式
    style_counts = {"Heading 1": 0, "Heading 2": 0, "Heading 3": 0}
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        if style_name in HEADING_STYLES:
            apply_paragraph_style(para, style_name)
            style_counts[style_name] += 1
    logger.info("✅ 标题样式已应用: %s", style_counts)

    # 3. 优化表格
    enhance_tables(doc)
    logger.info("✅ 已优化 %d 个表格", len(doc.tables))

    # 4. 页面设置
    set_page_setup(doc)
    logger.info("✅ 页面设置完成（A4 + 2.5/3.0cm 边距）")

    # 5. 页码
    add_page_number_footer(doc)
    logger.info("✅ 页码已添加")

    # 6. 保存
    doc.save(str(target))
    logger.info("✅ 模板已增强: %s", target)


def main() -> None:
    if sys.platform == "win32":
        try:
            sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
        except Exception:
            pass

    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s [%(levelname)s] %(message)s",
    )

    enhance_template()


if __name__ == "__main__":
    main()
