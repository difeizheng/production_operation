"""Tests for ReportGeneratorV2 (docxtpl-based) - Step 10"""
from __future__ import annotations

import sys
from pathlib import Path

import pytest

PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))


# ============================================================================
# 模板存在性测试
# ============================================================================

class TestTemplate:
    """模板存在性测试。"""

    def test_template_exists(self) -> None:
        """模板必须存在。"""
        from src.generator.report_generator_v2 import DEFAULT_TEMPLATE
        assert DEFAULT_TEMPLATE.exists(), (
            f"模板不存在: {DEFAULT_TEMPLATE}\n"
            "请先运行: python scripts/prepare_template.py"
        )

    def test_template_is_docx(self) -> None:
        """模板必须是 docx 格式。"""
        from src.generator.report_generator_v2 import DEFAULT_TEMPLATE
        from docx import Document
        doc = Document(str(DEFAULT_TEMPLATE))
        assert len(doc.paragraphs) > 0

    def test_template_has_jinja_placeholders(self) -> None:
        """模板应包含 Jinja2 占位符。"""
        from src.generator.report_generator_v2 import DEFAULT_TEMPLATE
        from docx import Document
        doc = Document(str(DEFAULT_TEMPLATE))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "{{" in all_text, "模板应包含 Jinja2 占位符"
        assert "v4_P6" in all_text or "v4_P" in all_text, "模板应包含原因占位符"


# ============================================================================
# 生成器测试
# ============================================================================

class TestReportGeneratorV2:
    """ReportGeneratorV2 测试。"""

    def test_init_with_default_template(self) -> None:
        """使用默认模板初始化。"""
        from src.generator.report_generator_v2 import ReportGeneratorV2
        gen = ReportGeneratorV2()
        assert gen.template_path.exists()

    def test_init_missing_template_raises(self) -> None:
        """模板不存在时抛错。"""
        from src.generator.report_generator_v2 import ReportGeneratorV2
        with pytest.raises(FileNotFoundError):
            ReportGeneratorV2(template_path=Path("/nonexistent/template.docx"))

    def test_generate_report_basic(
        self, tmp_path, sample_minimal_data  # type: ignore[no-untyped-def]
    ) -> None:
        """基本报告生成。"""
        from src.generator.report_generator_v2 import ReportGeneratorV2
        from docx import Document

        gen = ReportGeneratorV2(output_dir=str(tmp_path))
        output = gen.generate_report(
            data=sample_minimal_data, year=2026, week=21,
        )

        assert Path(output).exists()
        doc = Document(output)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        # 应包含年份周数
        assert "2026" in all_text or "21" in all_text

    def test_generate_with_reason_text(
        self, tmp_path, sample_minimal_data, sample_reason_segments  # type: ignore[no-untyped-def]
    ) -> None:
        """带 reason_text 的报告生成。"""
        from src.generator.report_generator_v2 import ReportGeneratorV2
        from docx import Document

        gen = ReportGeneratorV2(output_dir=str(tmp_path))
        output = gen.generate_report(
            data=sample_minimal_data, year=2026, week=21,
            reason_text=sample_reason_segments,
        )

        doc = Document(output)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        # 至少 P11 和 P12 应被渲染（其他占位符的匹配位置由模板决定）
        assert "测试原因文本 P11" in all_text
        assert "测试原因文本 P12" in all_text

    def test_generate_without_reason_uses_fallback(
        self, tmp_path, sample_minimal_data  # type: ignore[no-untyped-def]
    ) -> None:
        """无 reason_text 时使用 fallback（"（待补充）"）。"""
        from src.generator.report_generator_v2 import ReportGeneratorV2
        from docx import Document

        gen = ReportGeneratorV2(output_dir=str(tmp_path))
        output = gen.generate_report(
            data=sample_minimal_data, year=2026, week=21,
            reason_text=None,
        )

        doc = Document(output)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        # 至少有一个待补充（因为 reason_text=None）
        assert "（待补充）" in all_text or "待补充" in all_text


# ============================================================================
# Fixtures
# ============================================================================

@pytest.fixture
def sample_minimal_data() -> dict:
    """最小可用的 data 字典。"""
    return {
        "meta": {
            "year": 2026,
            "week": 21,
            "start_date": "5月16日",
            "end_date": "5月22日",
        },
        "domestic": {
            "electricity": {"total": 80.3, "hydro": 53.5},
            "price": {"total": 0.311},
        },
        "yoy": {
            "electricity": {"total": 0.033},
            "price": {"total": -0.009},
        },
        "wow": {
            "electricity": {"total": 0.157},
            "price": {"total": 0.001},
        },
        "revenue": {"total": 249000},
    }


@pytest.fixture
def sample_reason_segments() -> dict:
    """测试用 reason_text 字典。"""
    from src.generator.reason_resolver import ResolvedSegment

    def make_seg(placeholder: str, text: str) -> ResolvedSegment:
        return ResolvedSegment(
            placeholder=placeholder,
            final_text=text,
            automation_level="HIGH",
            source_slots=[],
            raw_text=text,
            polished=True,
            is_fallback=False,
        )

    return {
        "v4_P6_dom_elec_yoy_wow": make_seg("v4_P6_dom_elec_yoy_wow", "测试原因文本 P6"),
        "v4_P11_dom_price_yoy": make_seg("v4_P11_dom_price_yoy", "测试原因文本 P11"),
        "v4_P12_dom_price_wow": make_seg("v4_P12_dom_price_wow", "测试原因文本 P12"),
    }
