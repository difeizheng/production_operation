"""报告生成器 - 生成完整的周例会营销发言材料

生成的文档结构与真实文件一致：
    市场营销部汇报材料
    第XX周生产情况（X月XX日-X月XX日）

    一、上周销售情况
      （一）电量销售情况
        [概述段落]
        [水位表]
        1. 国内上网电价
          [同比分析段落]
          [环比分析段落]
          [各品类原因段落]
        [销售情况表]
        [发电收入总结段落]
        2. 国际上网电价（待补充）
        3. 市场化交易情况（待补充）
      （二）绿证、CCER核发交易情况（待补充）

    二、外部信息（待补充）

    三、重点工作情况（待补充）
    四、本周重点工作（待补充）
"""

import json
from pathlib import Path
from datetime import datetime
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.generator.table_builder import (
    create_sales_table,
    create_water_level_table,
    create_spot_price_table,
)
from src.generator.analysis_text import (
    generate_electricity_summary,
    generate_domestic_yoy_paragraph,
    generate_domestic_price_yoy_paragraph,
    generate_domestic_price_wow_paragraph,
    generate_category_price_reason_paragraph,
    generate_revenue_summary,
    generate_all_analysis_paragraphs,
    generate_spot_price_overview,
)


class ReportGenerator:
    """报告生成器，生成完整的周例会营销发言材料。"""

    def __init__(self, output_dir: str = "archive") -> None:
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def generate_report(
        self,
        data: Dict[str, Any],
        output_filename: Optional[str] = None,
        year: Optional[int] = None,
        week: Optional[int] = None,
    ) -> str:
        """生成完整的 Word 报告。

        Args:
            data: 由 AnalysisCollector.collect() 返回的数据
            output_filename: 输出文件名
            year: 年份
            week: 周数

        Returns:
            生成的文件路径
        """
        meta = data.get("meta", {})
        year = year or meta.get("year", datetime.now().year)
        week = week or meta.get("week", 1)

        if output_filename is None:
            output_filename = f"{year}年第{week}周周例会营销发言材料.docx"

        output_path = self.output_dir / str(year) / f"week{week}"
        output_path.mkdir(parents=True, exist_ok=True)
        output_file = output_path / output_filename

        # 创建文档
        doc = Document()

        # 设置默认字体
        style = doc.styles["Normal"]
        font = style.font
        font.name = "仿宋"
        font.size = Pt(14)

        # 构建文档
        self._build_title(doc, data, year, week)
        self._build_section_1(doc, data)
        self._build_section_2(doc, data)
        self._build_section_3(doc, data)
        self._build_section_4(doc, data)

        # 保存
        doc.save(output_file)
        return str(output_file)

    def _build_title(
        self, doc: Document, data: Dict[str, Any], year: int, week: int
    ) -> None:
        """构建标题部分。"""
        # 主标题：市场营销部汇报材料
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("市场营销部汇报材料")
        run.bold = True
        run.font.size = Pt(18)
        run.font.name = "黑体"

        # 副标题：第XX周生产情况
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta = data.get("meta", {})
        start_date = meta.get("start_date", "")
        end_date = meta.get("end_date", "")
        date_str = f"（{start_date}-{end_date}）" if start_date and end_date else ""

        heading = doc.add_heading(level=2)
        run = heading.add_run(f"第{week}周生产情况{date_str}")
        run.font.name = "黑体"

    def _build_section_1(self, doc: Document, data: Dict[str, Any]) -> None:
        """构建'一、上周销售情况'章节。"""
        doc.add_heading("一、上周销售情况", level=1)

        # （一）电量销售情况
        doc.add_paragraph("（一）电量销售情况")

        # 概述段落
        summary = generate_electricity_summary(data)
        doc.add_paragraph(summary)

        # 同比/环比分析段落（含原因）
        yoy_para = generate_domestic_yoy_paragraph(data)
        doc.add_paragraph(yoy_para)

        # 来水情况（仍需外部数据）
        doc.add_paragraph("（来水情况待补充）")

        # 水位表（占位，无数据时跳过）
        doc.add_paragraph("（水位表待补充）")

        # 1. 国内上网电价
        doc.add_heading("1.国内上网电价", level=3)

        # 同比分析段落
        price_yoy = generate_domestic_price_yoy_paragraph(data)
        doc.add_paragraph(price_yoy)

        # 各品类同比原因（占位）
        doc.add_paragraph("（各品类电价同比原因待补充）")

        # 环比分析段落
        price_wow = generate_domestic_price_wow_paragraph(data)
        doc.add_paragraph(price_wow)

        # 各品类环比原因
        for cat in ["水电", "风电", "光伏", "火电"]:
            reason = generate_category_price_reason_paragraph(cat, data, "wow")
            doc.add_paragraph(reason)

        # 表1：销售情况表
        doc.add_paragraph("表1 上周销售情况（亿千瓦时；元/千瓦时）")
        create_sales_table(doc, data)

        # 发电收入总结
        revenue_summary = generate_revenue_summary(data)
        doc.add_paragraph(revenue_summary)

        # 2. 国际上网电价（待补充）
        doc.add_heading("2.国际上网电价", level=3)
        intl = data.get("international", {})
        intl_price = intl.get("price", {}).get("total")
        intl_yoy = intl.get("yoy", {}).get("price_change")
        if intl_price is not None and intl_yoy is not None:
            doc.add_paragraph(
                f"第{data.get('meta', {}).get('week', '')}周，集团公司国际上网电价"
                f"约每千瓦时{intl_price:.3f}元，"
                f"同比度电{'提高' if intl_yoy > 0 else '下降'}"
                f"{abs(intl_yoy) * 100:.1f}分。"
                f"（详细分析待补充）"
            )
        else:
            doc.add_paragraph("（待补充）")

        # 3. 市场化交易情况（待补充）
        doc.add_heading("3.市场化交易情况", level=3)
        doc.add_paragraph("（待补充）")

        # （二）绿证、CCER
        doc.add_paragraph("（二）绿证、CCER核发交易情况")
        doc.add_paragraph("（待补充）")

    def _build_section_2(self, doc: Document, data: Dict[str, Any]) -> None:
        """构建'二、外部信息'章节。"""
        doc.add_heading("二、外部信息", level=1)
        doc.add_paragraph("（一）（待补充）")
        doc.add_paragraph("（二）（待补充）")

        # （三）现货市场价格信息
        spot_prices = data.get("spot_prices", {})
        regions = spot_prices.get("regions", [])

        if regions:
            # 有现货数据：生成概述段落 + 数据表
            overview = generate_spot_price_overview(data)
            doc.add_paragraph(f"（三）{overview}")
            doc.add_paragraph("表2 上周现货市场均价（元/千瓦时）")

            # 构建表格数据
            spot_data = spot_prices.get("data", {})
            price_data = {}
            for region in regions:
                d = spot_data.get(region, {})
                price_data[region] = {
                    "avg": d.get("avg"),
                    "yoy": d.get("yoy"),
                    "wow": d.get("wow"),
                }
            create_spot_price_table(doc, regions=regions, price_data=price_data)
        else:
            # 无现货数据：占位
            doc.add_paragraph("（三）现货市场价格信息。")
            doc.add_paragraph("表2 上周现货市场均价（元/千瓦时）")
            create_spot_price_table(doc)

    def _build_section_3(self, doc: Document, data: Dict[str, Any]) -> None:
        """构建'三、重点工作情况'章节。"""
        doc.add_heading("三、重点工作情况", level=1)
        doc.add_paragraph("（待补充）")

    def _build_section_4(self, doc: Document, data: Dict[str, Any]) -> None:
        """构建'四、本周重点工作'章节。"""
        p = doc.add_paragraph()
        run = p.add_run("四、本周重点工作")
        run.bold = True
        doc.add_paragraph("（待补充）")


def main() -> None:
    """命令行入口"""
    import argparse

    parser = argparse.ArgumentParser(description="生成周报")
    parser.add_argument("--input", required=True, help="JSON 数据文件路径")
    parser.add_argument("--output", help="输出文件名")
    parser.add_argument("--year", type=int, help="年份")
    parser.add_argument("--week", type=int, help="周数")

    args = parser.parse_args()

    # 加载数据
    with open(args.input, "r", encoding="utf-8") as f:
        data = json.load(f)

    # 生成报告
    generator = ReportGenerator()
    path = generator.generate_report(
        data, output_filename=args.output, year=args.year, week=args.week
    )
    print(f"报告已生成: {path}")


if __name__ == "__main__":
    main()
