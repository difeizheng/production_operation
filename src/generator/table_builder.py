"""Word 表格构建器 - 创建和填充 Word 文档中的表格

支持三种表格：
    - 销售情况表（Table 1）：10行×7列
    - 水位表（Table 0）：4行×6列
    - 现货市场均价表（Table 2）：4行×N列
"""

from typing import Any, Dict, List, Optional

from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def create_sales_table(doc: Any, data: Dict[str, Any]) -> Any:
    """创建销售情况表（10行×7列）。

    结构与真实文档 Table 1 一致：
        Row 0: 表头 ['', '水电', '新能源', '风电', '光伏', '火电', '合计']
        Row 1: 国内上网电量
        Row 2: 同比
        Row 3: 环比
        Row 4: 国内上网电价
        Row 5: 同比
        Row 6: 环比
        Row 7: 国内发电收入
        Row 8: 同比
        Row 9: 环比

    Args:
        doc: python-docx Document 对象
        data: 含 report_table_1 的数据字典

    Returns:
        创建的 Table 对象
    """
    table_data = data.get("report_table_1", {})
    headers = table_data.get("headers", ["", "水电", "新能源", "风电", "光伏", "火电", "合计"])
    row_labels = table_data.get("row_labels", [
        "国内上网电量", "同比", "环比",
        "国内上网电价", "同比", "环比",
        "国内发电收入", "同比", "环比",
    ])
    rows_data = table_data.get("data", [[] for _ in range(9)])

    # 创建表格 (10行 x 7列)
    table = doc.add_table(rows=10, cols=7)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 填充表头
    for col_idx, header in enumerate(headers):
        cell = table.rows[0].cells[col_idx]
        cell.text = header
        _set_cell_font(cell, bold=True)

    # 填充数据行
    for row_idx in range(min(9, len(rows_data))):
        row = table.rows[row_idx + 1]

        # 第一列：行标签
        label = row_labels[row_idx] if row_idx < len(row_labels) else ""
        row.cells[0].text = label

        # 数据列
        row_values = rows_data[row_idx] if row_idx < len(rows_data) else []
        col_keys = ["hydro", "new_energy", "wind", "solar", "thermal", "total"]

        for col_idx in range(6):
            if col_idx < len(row_values):
                value = row_values[col_idx]
                row.cells[col_idx + 1].text = _format_table_value(
                    value, row_idx, col_idx
                )

    return table


def create_water_level_table(
    doc: Any,
    stations: Optional[List[str]] = None,
    dates: Optional[List[str]] = None,
    levels: Optional[Dict[str, List[Optional[float]]]] = None,
) -> Any:
    """创建水位表（4行×6列）。

    结构：
        Row 0: ['水位（米）', '三峡', '向家坝', '溪洛渡', '白鹤滩', '乌东德']
        Row 1: [日期1, 值...]
        Row 2: [日期2, 值...]
        Row 3: ['环比', 差值...]

    Args:
        doc: python-docx Document 对象
        stations: 站名列表
        dates: 两个日期字符串
        levels: {站名: [日期1水位, 日期2水位]}

    Returns:
        创建的 Table 对象
    """
    if stations is None:
        stations = ["三峡", "向家坝", "溪洛渡", "白鹤滩", "乌东德"]
    if dates is None:
        dates = ["x月xx日", "x月xx日"]
    if levels is None:
        levels = {s: [None, None] for s in stations}

    table = doc.add_table(rows=4, cols=6)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    headers = ["水位（米）"] + stations
    for col_idx, header in enumerate(headers):
        cell = table.rows[0].cells[col_idx]
        cell.text = header
        _set_cell_font(cell, bold=True)

    # 日期1行
    table.rows[1].cells[0].text = dates[0]
    for i, s in enumerate(stations):
        val = levels.get(s, [None, None])[0]
        table.rows[1].cells[i + 1].text = f"{val:.2f}" if val is not None else ""

    # 日期2行
    table.rows[2].cells[0].text = dates[1]
    for i, s in enumerate(stations):
        val = levels.get(s, [None, None])[1]
        table.rows[2].cells[i + 1].text = f"{val:.2f}" if val is not None else ""

    # 环比行
    table.rows[3].cells[0].text = "环比"
    for i, s in enumerate(stations):
        vals = levels.get(s, [None, None])
        if vals[0] is not None and vals[1] is not None:
            diff = vals[1] - vals[0]
            table.rows[3].cells[i + 1].text = f"{diff:+.2f}"
        else:
            table.rows[3].cells[i + 1].text = ""

    return table


def create_spot_price_table(
    doc: Any,
    regions: Optional[List[str]] = None,
    price_data: Optional[Dict[str, Dict[str, Optional[float]]]] = None,
) -> Any:
    """创建现货市场均价表（4行×N列）。

    结构：
        Row 0: ['地区', '广东', '山西', '山东', ...]
        Row 1: ['均价', 值...]
        Row 2: ['同比', 值...]
        Row 3: ['环比', 值...]

    Args:
        doc: python-docx Document 对象
        regions: 地区列表
        price_data: {地区: {avg, yoy, wow}}

    Returns:
        创建的 Table 对象
    """
    if regions is None:
        regions = ["广东", "山西", "山东", "甘肃", "蒙西", "湖北", "浙江", "陕西"]
    if price_data is None:
        price_data = {r: {"avg": None, "yoy": None, "wow": None} for r in regions}

    col_count = 1 + len(regions)
    table = doc.add_table(rows=4, cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    table.rows[0].cells[0].text = "地区"
    _set_cell_font(table.rows[0].cells[0], bold=True)
    for i, region in enumerate(regions):
        table.rows[0].cells[i + 1].text = region
        _set_cell_font(table.rows[0].cells[i + 1], bold=True)

    # 均价行
    row_labels = ["均价", "同比", "环比"]
    data_keys = ["avg", "yoy", "wow"]
    for row_idx in range(3):
        table.rows[row_idx + 1].cells[0].text = row_labels[row_idx]
        for i, region in enumerate(regions):
            val = price_data.get(region, {}).get(data_keys[row_idx])
            if val is not None:
                if row_idx == 0:
                    table.rows[row_idx + 1].cells[i + 1].text = f"{val:.3f}"
                else:
                    table.rows[row_idx + 1].cells[i + 1].text = f"{val * 100:.1f}%"
            else:
                table.rows[row_idx + 1].cells[i + 1].text = ""

    return table


# ============================================================================
# 内部工具函数
# ============================================================================


def _set_cell_font(cell: Any, bold: bool = False, size: int = 9) -> None:
    """设置单元格字体。

    Args:
        cell: 表格单元格
        bold: 是否粗体
        size: 字号（磅）
    """
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(size)
            run.font.bold = bold


def _format_table_value(
    value: Optional[float], row_idx: int, col_idx: int
) -> str:
    """根据行列位置格式化表格值。

    行索引决定格式化方式：
        0, 3, 6 (电量/电价/收入绝对值)
        1, 4, 7 (同比率)
        2, 5, 8 (环比率)

    Args:
        value: 原始值
        row_idx: 数据行索引（0-8）
        col_idx: 列索引（0-5）

    Returns:
        格式化后的字符串
    """
    if value is None:
        return ""

    # 同比/环比行（百分比）
    if row_idx in (1, 2, 4, 5, 7, 8):
        return f"{value * 100:.1f}%"

    # 电价行（3位小数）
    if row_idx in (3,):
        return f"{value:.3f}"

    # 电量/收入行（1位小数）
    return f"{value:.1f}"
