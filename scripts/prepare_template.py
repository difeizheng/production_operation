"""模板准备脚本 - 在 V3 模板上插入 Jinja2 占位符

输入：data/templates/report_template.docx（V3 复制）
输出：data/templates/report_template_jinja.docx（含 {{ 占位符 }}）

占位符布局（与 report_generator 集成）：
    {{ title }} - 主标题
    {{ subtitle }} - 副标题
    {{ year }}, {{ week }} - 年周
    {{ date_range }} - 日期范围
    {{ dom_overview }} - 国内电量概览
    {{ v4_P6_dom_elec_yoy_wow }} - 原因段落 6
    {{ v4_P11_dom_price_yoy }} - 原因段落 11
    ... 等等

使用：
    PYTHONPATH=. python scripts/prepare_template.py
"""
from __future__ import annotations

import io
import logging
import re
import sys
from pathlib import Path
from typing import Dict, List

PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

logger = logging.getLogger(__name__)


# ============================================================================
# 占位符映射
# ============================================================================

# 段落替换规则：(匹配模式, 替换内容)
# 匹配模式: 段落中包含的文本
REPLACEMENTS: List[Dict[str, str]] = [
    # 标题
    {"find": "市场营销部汇报材料", "replace": "{{ title|default('市场营销部汇报材料') }}"},
    # 副标题（如 "第21周生产情况"）
    {"find": "第21周生产情况", "replace": "第{{ week|default('XX') }}周生产情况"},
    # 国内电量概览
    {
        "find": "上周，集团公司合计上网电量",
        "replace": "{{ dom_overview }} 上周，集团公司合计上网电量",
    },
    # V4 P6 - 国内电量同比/环比原因
    {
        "find": "上周，集团公司国内上网电量80.3",
        "replace": "上周，集团公司国内上网电量80.3",
        "after": "{{ v4_P6_dom_elec_yoy_wow }}",
    },
    # V4 P11 - 国内电价同比 品类原因
    {
        "find": "上周，集团公司国内平均上网电价约每千瓦时0.311元",
        "after": "{{ v4_P11_dom_price_yoy }}",
    },
    # V4 P12 - 国内电价环比 品类原因
    {
        "find": "国内平均上网电价环比度电提高",
        "after": "{{ v4_P12_dom_price_wow }}",
    },
    # V4 P13-P16 - 品类级环比原因
    {
        "find": "水电电价环比下降的原因：",
        "replace": "水电电价环比下降的原因：{{ v4_P13_hydro_price_wow }}",
    },
    {
        "find": "风电电价环比提高的原因：",
        "replace": "风电电价环比提高的原因：{{ v4_P14_wind_price_wow }}",
    },
    {
        "find": "光伏电价环比提高的原因：",
        "replace": "光伏电价环比提高的原因：{{ v4_P15_solar_price_wow }}",
    },
    {
        "find": "火电电价环比提高原因：",
        "replace": "火电电价环比提高原因：{{ v4_P16_thermal_price_wow }}",
    },
    # V4 P18 - 发电收入
    {
        "find": "上周，集团公司国内发电收入",
        "after": "{{ v4_P18_dom_revenue }}",
    },
    # V4 P20/P21 - 国际电价
    {
        "find": "第21周，集团公司国际上网电价",
        "after": "{{ v4_P20_intl_price_yoy }}",
    },
    # V4 P23 - 市场化水电
    {
        "find": "水电项目：长江干流梯级电站",
        "after": "{{ v4_P23_market_hydro }}",
    },
    # V4 P24 - 市场化新能源
    {
        "find": "新能源项目（参与28个省",
        "after": "{{ v4_P24_market_new_energy }}",
    },
    # V4 P25 - 市场化火电
    {
        "find": "火电项目目前仅燃煤机组",
        "after": "{{ v4_P25_market_thermal }}",
    },
    # V4 P27 - 绿证
    {
        "find": "上周，集团公司核发绿证",
        "after": "{{ v4_P27_green_cert }}",
    },
]


def prepare_template(
    source: Path = PROJECT_ROOT / "data" / "templates" / "report_template.docx",
    target: Path = PROJECT_ROOT / "data" / "templates" / "report_template_jinja.docx",
) -> None:
    """在 V3 模板上插入 Jinja2 占位符。

    通过解析 docx 的 XML 直接修改文本。
    """
    import zipfile
    import shutil

    if not source.exists():
        logger.error("源模板不存在: %s", source)
        return

    logger.info("源: %s", source)
    logger.info("目标: %s", target)

    # 复制源到目标
    shutil.copy2(str(source), str(target))

    # 用 python-docx 打开
    from docx import Document
    doc = Document(str(target))

    modified_count = 0
    for para in doc.paragraphs:
        text = para.text
        if not text:
            continue

        original = text

        for rule in REPLACEMENTS:
            find = rule["find"]
            if find not in text:
                continue

            # 'replace' 模式
            if "replace" in rule:
                text = text.replace(find, rule["replace"])
            # 'after' 模式（在段尾追加）
            elif "after" in rule:
                # 先去掉 find 之后的所有内容（如果有的话）
                # 然后附加新内容
                base = text.split(find)[0] + find
                text = base + rule["after"]

        if text != original:
            # 修改段落文本（保留第一个 run 的格式）
            if para.runs:
                # 清空所有 run 内容，把新文本放第一个 run
                first_run = para.runs[0]
                first_run.text = text
                for run in para.runs[1:]:
                    run.text = ""
                modified_count += 1
            else:
                # 没有 run，直接添加
                para.add_run(text)
                modified_count += 1

    doc.save(str(target))
    logger.info("✅ 修改了 %d 个段落", modified_count)


def main() -> None:
    if sys.platform == "win32":
        try:
            sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
        except Exception:
            pass

    logging.basicConfig(level=logging.INFO, format="%(message)s")
    prepare_template()


if __name__ == "__main__":
    main()
