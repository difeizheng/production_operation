"""Unit tests for table template preparation (Step 13)"""
from __future__ import annotations

import sys
from pathlib import Path

import pytest

PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))


class TestBuildSalesRows:
    """build_sales_rows 测试。"""

    def test_full_data(self) -> None:
        from scripts.prepare_table_templates import build_sales_rows

        data = {
            "report.electricity.hydro": 59.045,
            "report.electricity.new_energy": 17.957,
            "report.electricity.wind": 10.363,
            "report.electricity.solar": 7.594,
            "report.electricity.thermal": 3.118,
            "report.electricity.total": 80.276,
            "report.yoy_electricity.hydro": 0.066,
            "report.yoy_electricity.total": 0.033,
            "report.price.hydro": 0.283,
            "report.wow_price.hydro": -0.0044,
        }
        rows = build_sales_rows(data)

        assert "row_1" in rows  # 销量
        assert rows["row_1"]["hydro"] == "59.05"
        assert rows["row_1"]["total"] == "80.28"

        assert "row_2" in rows  # 同比
        assert rows["row_2"]["hydro"] == "+6.6%"
        assert rows["row_2"]["total"] == "+3.3%"

        assert "row_5" in rows  # 电价同比
        # yoy_price 不在测试数据中，应为 —
        assert rows["row_5"]["hydro"] == "—"

    def test_missing_data(self) -> None:
        from scripts.prepare_table_templates import build_sales_rows

        rows = build_sales_rows({})
        for row_idx in range(1, 10):
            key = f"row_{row_idx}"
            assert key in rows
            # 所有值应为 —
            for col in ["hydro", "wind", "total"]:
                assert rows[key][col] == "—"


class TestBuildSpotRows:
    """build_spot_rows 测试。"""

    def test_spot_data(self) -> None:
        from scripts.prepare_table_templates import build_spot_rows

        data = {
            "spot_prices": {
                "data": {
                    "广东": {"avg": 0.366, "yoy": -0.014, "wow": 0.05},
                    "山东": {"avg": 0.333, "yoy": 0.228, "wow": 0.031},
                }
            }
        }
        spot = build_spot_rows(data)

        assert "spot_avg" in spot
        assert "spot_yoy" in spot
        assert "spot_wow" in spot

        assert spot["spot_avg"]["广东"] == "0.366"
        assert spot["spot_yoy"]["广东"] == "-1.4%"
        assert spot["spot_wow"]["山东"] == "+3.1%"

    def test_empty_spot_data(self) -> None:
        from scripts.prepare_table_templates import build_spot_rows

        spot = build_spot_rows({})
        assert spot["spot_avg"] == {}
        assert spot["spot_yoy"] == {}


class TestFormatValue:
    """_format_value 测试。"""

    def test_pct_format(self) -> None:
        from scripts.prepare_table_templates import _format_value

        assert _format_value(0.033, "pct") == "+3.3%"
        assert _format_value(-0.014, "pct") == "-1.4%"
        assert _format_value(0, "pct") == "0.0%"
        assert _format_value(None, "pct") == "—"

    def test_fen_format(self) -> None:
        from scripts.prepare_table_templates import _format_value

        assert _format_value(-0.0044, "fen") == "-0.4"
        assert _format_value(0.011, "fen") == "+1.1"

    def test_yuan_per_kwh_format(self) -> None:
        from scripts.prepare_table_templates import _format_value

        assert _format_value(0.366, "yuan_per_kwh") == "0.366"
        assert _format_value(0.283015, "yuan_per_kwh") == "0.283"

    def test_number_format(self) -> None:
        from scripts.prepare_table_templates import _format_value

        assert _format_value(59.045, "number") == "59.05"
        assert _format_value(80.276, "number") == "80.28"


# ============================================================================
# 端到端测试
# ============================================================================

class TestEndToEndTableRender:
    """表格端到端渲染测试。"""

    def test_render_with_tables(
        self, tmp_path, sample_full_data  # type: ignore[no-untyped-def]
    ) -> None:
        """带表格数据的端到端渲染。"""
        from src.generator.report_generator_v2 import ReportGeneratorV2
        from docx import Document

        gen = ReportGeneratorV2(output_dir=str(tmp_path))
        output = gen.generate_report(
            data=sample_full_data, year=2026, week=21,
        )

        doc = Document(output)
        # 应能找到表 1（销售情况表）
        assert len(doc.tables) >= 1

        # 检查销售表的数据
        sales_table = None
        for t in doc.tables:
            if len(t.rows) >= 10 and len(t.columns) >= 6:
                sales_table = t
                break
        assert sales_table is not None

        # 第一行（表头）应有 "水电"
        header_row = sales_table.rows[0]
        assert "水电" in header_row.cells[1].text

        # 第 2 行（销量行）的第 1 列应为数字（59.05）
        sales_row = sales_table.rows[1]
        cell_text = sales_row.cells[1].text
        assert "59" in cell_text  # 59.05


# ============================================================================
# Fixtures
# ============================================================================

@pytest.fixture
def sample_full_data() -> dict:
    """完整 data 字典（含 report 字段）。"""
    return {
        "meta": {"year": 2026, "week": 21, "start_date": "5月16日", "end_date": "5月22日"},
        "domestic": {
            "electricity": {"total": 80.276, "hydro": 59.045, "new_energy": 17.957, "wind": 10.363, "solar": 7.594, "thermal": 3.118},
            "price": {"total": 0.311, "hydro": 0.283},
        },
        "yoy": {"electricity": {"total": 0.033, "hydro": 0.066}, "price": {"total": -0.009, "hydro": -0.001}},
        "wow": {"electricity": {"total": 0.157, "hydro": 0.192}, "price": {"total": 0.001, "hydro": -0.0044}},
        "revenue": {"total": 249000, "hydro": 16710},
        # report.* 字段（用于 build_sales_rows）
        "report.electricity.hydro": 59.045,
        "report.electricity.new_energy": 17.957,
        "report.electricity.wind": 10.363,
        "report.electricity.solar": 7.594,
        "report.electricity.thermal": 3.118,
        "report.electricity.total": 80.276,
        "report.yoy_electricity.hydro": 0.066,
        "report.yoy_electricity.new_energy": 0.046,
        "report.yoy_electricity.wind": 0.118,
        "report.yoy_electricity.solar": -0.038,
        "report.yoy_electricity.thermal": -0.376,
        "report.yoy_electricity.total": 0.033,
        "report.wow_electricity.hydro": 0.192,
        "report.wow_electricity.total": 0.157,
        "report.price.hydro": 0.283,
        "report.price.total": 0.311,
        "report.wow_price.hydro": -0.0044,
        # 现货
        "spot_prices": {
            "data": {
                "广东": {"avg": 0.366, "yoy": -0.014, "wow": 0.05},
                "山东": {"avg": 0.333, "yoy": 0.228, "wow": 0.031},
            }
        },
    }
